import os, re, json, docx
from datetime import datetime
import easygui
import logging
import shutil
from docx.oxml.ns import qn

logging.basicConfig(level=logging.INFO, format="%(levelname)s - %(message)s")
logger = logging.getLogger(__name__)


def getFilePath(folderPath, fileName, fileType):
    if not os.path.exists(folderPath):
        os.makedirs(folderPath)
    targetFilePath = fileName + "." + fileType
    return os.path.join(folderPath, targetFilePath)


def getDirPath(folderPath, targetDirPath):
    if not os.path.exists(folderPath):
        os.makedirs(folderPath)
    return os.path.join(folderPath, targetDirPath)


class TextJsonConverter(object):
    def __init__(
        self,
        sourcePath,
        targetFolderPath="",
        datePattern="\d{4}-\d{2}-\d{2}",
        timePattern="\d{2}:\d{2}\Z",
        maxNameLength=40,
        type="Android",
    ):
        self.sourcePath = sourcePath
        self.bookTitle = sourcePath.split("\\")[-1].split(".")[0]
        if targetFolderPath == "":
            self.targetFolderPath = self.bookTitle
        else:
            self.targetFolderPath = targetFolderPath

        self.TJC_Dict = {}
        self.TJC_Dict["基本信息"] = {
            "书名": self.bookTitle,
            "开始日期": 0,
            "结束日期": 0,
            "处理日期": datetime.now().strftime("%Y-%m-%d"),
        }

        self.datePattern = datePattern  # re pattern for date recognition
        self.timePattern = timePattern  # re pattern for time recogization
        self.maxNameLength = maxNameLength  # the longest supported name
        logger.info("TextJsonConverter created!")

    def addContent(self, name, date, content):
        if name in self.TJC_Dict:
            if date in self.TJC_Dict[name]:
                self.TJC_Dict[name][date].append("\n".join(content))
            else:
                self.TJC_Dict[name][date] = ["\n".join(content)]
        else:
            self.TJC_Dict[name] = {date: ["\n".join(content)]}

    def checkStatus(self):
        if self.sourcePath == "":
            return False
        elif self.targetFolderPath == "":
            return False
        return True

    def handleTxt_Android(self, txt):
        name, date = "temp", "temp"
        last = 0

        for i in range(len(txt)):
            re_date = re.search(self.datePattern, txt[i])
            if re_date:
                date = re_date.group(0)
                content = txt[last : i - 1]
                self.addContent(name, date, content)
                # logger.info("date=" + date + ", name=" + name)

                last = i + 1
                if (self.TJC_Dict["基本信息"]["开始日期"] == 0) and (date != "temp"):
                    logger.info(date)
                    self.TJC_Dict["基本信息"]["开始日期"] = date
                self.TJC_Dict["基本信息"]["结束日期"] = date

            re_name = re.search(self.timePattern, txt[i])
            if re_name and (len(txt[i]) < self.maxNameLength):
                time = txt[i][-5:]

                content = txt[last : i - 1]
                self.addContent(name, date, content)
                logger.info("date=" + date + ", name=" + name)

                if txt[i][-6] != chr(32):
                    name = txt[i][:-5]
                else:
                    name = txt[i][:-6]

                last = i + 1

        content = txt[last:]
        self.addContent(name, date, content)
        del self.TJC_Dict["temp"]
        return True

    def handleTxt_IOS(self, txt):
        name, date = "temp", "temp"
        last = 0

        for i in range(len(txt)):
            re_date = re.search(self.datePattern, txt[i])
            if re_date:

                newDate = re_date.group(0)

                if len(txt[i - 1]) > 1:
                    newName = txt[i - 1][:-1]
                else:
                    newName = txt[i - 2][:-1]

                while newName[-1] in [":", "*"]:
                    newName = newName[:-1]

                content = txt[last : i - 1]
                self.addContent(name, date, content)
                # logger.info("date=" + date + ", name=" + name)

                last = i + 1
                if (self.TJC_Dict["基本信息"]["开始日期"] == 0) and (date != "temp"):
                    self.TJC_Dict["基本信息"]["开始日期"] = date
                self.TJC_Dict["基本信息"]["结束日期"] = date
                name, date = newName, newDate

        content = txt[last:]
        self.addContent(name, date, content)
        # logger.info("date=" + date + ", name=" + name)

        del self.TJC_Dict["temp"]
        return True

    def generateJsonFromTxt(self, platform="Android"):
        if not self.checkStatus():
            return False

        txt = open(self.sourcePath, encoding="utf-8").read().split("\n")

        if platform == "Android":
            self.handleTxt_Android(txt)
        elif platform == "IOS":
            self.handleTxt_IOS(txt)
        else:
            logger.error("Platform unsupported!")
            return False

        jsonName = getFilePath(self.targetFolderPath, self.bookTitle, "json")
        json.dump(
            self.TJC_Dict, open(jsonName, "w", encoding="utf-8"), ensure_ascii=False
        )
        logger.info("JSON file generated successfully!")
        return jsonName


class JsonTxtConverter(object):
    def __init__(self, sourcePath, targetFolderPath="", threshold=1):
        self.sourcePath = sourcePath

        self.JTC_Dict = json.load(open(self.sourcePath, "rb"))
        self.bookTitle = self.JTC_Dict["基本信息"]["书名"]

        if targetFolderPath == "":
            self.targetFolderPath = self.bookTitle
        else:
            self.targetFolderPath = targetFolderPath

        self.threshold = threshold
        logger.info("JsonTxtConverter created!")

    def handleTxt_HP(self, postfix="读书感想合集"):
        targetPath_HP = self.bookTitle + "_" + postfix
        targetPath_HP = getFilePath(self.targetFolderPath, targetPath_HP, "txt")
        fileTxt_HP = open(targetPath_HP, "w", encoding="utf-8")

        for name in self.JTC_Dict.keys():
            if name == "基本信息":
                fileTxt_HP.write(name + ":\n")
                for (k, v) in self.JTC_Dict[name].items():
                    fileTxt_HP.write("【" + k + "】")
                    fileTxt_HP.write(v + "\n")
                continue

            if self.JTC_Dict[name].__len__() < self.threshold:
                continue

            name_print = "\n" * 5 + "参与者：" + name
            fileTxt_HP.write(name_print)

            for date in self.JTC_Dict[name].keys():
                date_print = "\n【" + date + "】\n"
                fileTxt_HP.write(date_print)

                NoteList = self.JTC_Dict[name][date]
                for i in range(len(NoteList)):
                    content = NoteList[i]
                    fileTxt_HP.write(content)

        fileTxt_HP.close()
        logger.info("Txt HodgePodge file generated into: " + targetPath_HP)

    def handleTxt_SP(self):
        for name in self.JTC_Dict.keys():
            if name == "基本信息":
                continue

            if self.JTC_Dict[name].__len__() < self.threshold:
                continue

            target_SP = self.bookTitle + "_" + name
            target_SP = getFilePath(self.targetFolderPath, target_SP, "txt")
            file_SP = open(target_SP, "w", encoding="utf-8")

            for date in self.JTC_Dict[name].keys():
                date_print = "\n【" + date + "】\n"
                file_SP.write(date_print)

                NoteList = self.JTC_Dict[name][date]
                for i in range(len(NoteList)):
                    content = NoteList[i]
                    file_SP.write(content)
            file_SP.close()
        logger.info("Word HodgePodge file generated into: " + self.targetFolderPath)

    def handleWord_HP(self, postfix="读书感想合集"):
        targetWord_HP = self.bookTitle + "_" + postfix
        targetWord_HP = getFilePath(self.targetFolderPath, targetWord_HP, "docx")
        fileWord_HP = docx.Document()
        fileWord_HP.styles["Normal"].font.name = u"宋体"
        fileWord_HP.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), u"宋体")

        for name in self.JTC_Dict.keys():
            if name == "基本信息":
                for (k, v) in self.JTC_Dict[name].items():
                    fileWord_HP.add_paragraph("【" + k + "】 " + v)
                fileWord_HP.add_paragraph("")
                continue

            if self.JTC_Dict[name].__len__() < self.threshold:
                continue

            name_print = "参与者：" + name
            fileWord_HP.add_paragraph(name_print)

            for date in self.JTC_Dict[name].keys():

                date_print = "【" + date + "】"
                fileWord_HP.add_paragraph("")
                fileWord_HP.add_paragraph(date_print)

                NoteList = self.JTC_Dict[name][date]
                for i in range(len(NoteList)):
                    content = NoteList[i]
                    fileWord_HP.add_paragraph(content)

            fileWord_HP.add_page_break()

        fileWord_HP.save(targetWord_HP)
        logger.info("Word HodgePodge file generated into: " + targetWord_HP)
        return targetWord_HP

    def handleWord_SP(self):

        for name in self.JTC_Dict.keys():
            if name == "基本信息":
                continue

            if self.JTC_Dict[name].__len__() < self.threshold:
                continue

            fileWord_SP = docx.Document()
            fileWord_SP.styles["Normal"].font.name = u"宋体"
            fileWord_SP.styles["Normal"]._element.rPr.rFonts.set(
                qn("w:eastAsia"), u"宋体"
            )

            for date in self.JTC_Dict[name].keys():

                date_print = "【" + date + "】"
                fileWord_SP.add_paragraph(date_print)

                NoteList = self.JTC_Dict[name][date]
                for i in range(len(NoteList)):
                    content = NoteList[i]
                    fileWord_SP.add_paragraph(content)

            fileWord_SP.add_page_break()

            targetWord_SP = getFilePath(self.targetFolderPath, name, "docx")
            fileWord_SP.save(targetWord_SP)

        logger.info("Word Separate file generated into: " + self.targetFolderPath)
        return True

    def saveTxtFromJson(self, targetType="txt"):

        if targetType == "txt":
            self.handleTxt_HP()
            self.handleTxt_SP()

        elif targetType == "word":
            self.handleWord_HP()
            self.handleWord_SP()
        else:
            logger.error("Target file type not supported!")

        return True

    def clean(self, cleanSourseJson=True, cleanFolder=False):
        if cleanFolder == True:
            shutil.rmtree(self.targetFolderPath, True)
            logger.warning("cleaned folder: " + self.targetFolderPath)
        elif cleanSourseJson == True:
            # todo
            print(self.sourcePath)
            # os.remove(self.sourcePath)
            logger.warning("cleaned source: " + self.sourcePath)


if __name__ == "__main__":

    minDiarys = 1

    pattern1 = "\d{4}-\d{2}-\d{2}"
    pattern2 = "\d{2}\.\d{2}"
    pattern3 = "[＃#]\d{2}\.\d{2}[＃#]"

    sourcePath = easygui.fileopenbox(msg="选择源文件", title="Source File")
    platform = easygui.choicebox(
        msg="选择聊天记录的来源", title="IOS or Android", choices=("IOS", "Android")
    )
    datePattern = easygui.choicebox(
        msg="正则表达式的日期形式",
        title="Regular Expression Pattern",
        choices=(pattern1, pattern2, pattern3),
    )

    if sourcePath == None or platform == None or datePattern == None:
        logger.error("Invalid Input. Exit.")
        exit()
    

    TJC = TextJsonConverter(sourcePath, datePattern=datePattern)
    jsonFileName = TJC.generateJsonFromTxt(platform=platform)
    logger.info("JSON file generated: " + jsonFileName)

    JTC = JsonTxtConverter(jsonFileName, threshold=minDiarys)
    JTC.clean(cleanSourseJson=False, cleanFolder=True)
    JTC.saveTxtFromJson(targetType="word")
    # JTC.saveTxtFromJson(targetType="txt")
